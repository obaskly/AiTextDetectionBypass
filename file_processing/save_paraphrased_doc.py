from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate
from pypdf import PdfReader, PdfWriter
from io import BytesIO
import os 

def get_output_filename(input_path):
    base_dir = os.path.dirname(input_path)
    ext = os.path.splitext(input_path)[1]
    return os.path.join(base_dir, f"paraphrased{ext}")

def save_as_docx(input_path, paraphrased_text, save_path=None):
    try:
        output_path = save_path or get_output_filename(input_path)  # Use custom path if provided
        
        if os.path.exists(output_path):
            doc = Document(output_path)
        else:
            doc = Document()
        
        doc.add_paragraph(paraphrased_text)
        doc.save(output_path)
        print(f"Successfully saved paraphrased text to {output_path}")
    except Exception as e:
        raise ValueError(f"Error saving DOCX: {str(e)}")

def save_as_txt(input_path, paraphrased_text, save_path=None):
    try:
        output_path = save_path or get_output_filename(input_path)  # Use custom path if provided

        with open(output_path, 'a', encoding='utf-8') as f:
            f.write(paraphrased_text + '\n') 
        
        print(f"Successfully saved paraphrased text to {output_path}")
    except Exception as e:
        raise ValueError(f"Error saving TXT: {str(e)}")

def save_as_pdf(input_path, paraphrased_text, save_path=None):
    try:
        output_path = save_path or get_output_filename(input_path)  # Use custom path if provided

        existing_text = ""
        if os.path.exists(output_path):
            reader = PdfReader(output_path)
            for page in reader.pages:
                existing_text += page.extract_text() + "\n"

        combined_text = existing_text.strip() + "\n\n" + paraphrased_text.strip()

        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        style = ParagraphStyle(
            'Normal',
            fontSize=12,
            leading=16,
            spaceBefore=12,
            spaceAfter=12,
            firstLineIndent=24
        )

        story = []
        for paragraph in combined_text.split("\n\n"):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), style))

        doc.build(story)

        print(f"Successfully saved paraphrased text to {output_path}")
    except Exception as e:
        raise ValueError(f"Error saving PDF: {str(e)}")
